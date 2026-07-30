"""python-docx renderer for the CD&A report.

Consumes a :class:`~app.services.cda.schema.CDADataset` plus a flat list of
content blocks (see ``schema`` for the block DSL) and returns raw ``.docx``
bytes. No data access or LLM calls happen here — everything was resolved
upstream.

Design mirrors ``pdf_builder`` one-for-one (same palette, same block types)
so the Word document reads like the PDF did: a cover, a running header and a
running footer ("<proxy_year> Proxy Statement · genpact"), coral section
rules, zebra tables, and the cream/amber "Key Financial Highlights" panel —
here a single continuous cream block rather than separate cards.
"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from .schema import Block, CDADataset

# --- palette (same hex as pdf_builder) -------------------------------------
_INK = RGBColor(0x16, 0x16, 0x16)
_MUTED = RGBColor(0x6E, 0x6E, 0x6E)
_ACCENT = RGBColor(0xC0, 0x39, 0x2B)      # deep coral/red — section rules
_ACCENT_HEX = "C0392B"                      # same, as a shading/border hex
_ACCENT_DK = RGBColor(0x78, 0x23, 0x1B)
_AMBER = RGBColor(0xF4, 0xA5, 0x22)       # highlight-card headline values
_CREAM = "FBF4E6"                          # highlight-panel fill (shading hex)
_LINE = "D2CDC3"                           # table border
_ZEBRA = "F7F3EC"                          # alternate table row fill
_HEAD_BG = "F3E9D6"                        # table header fill

_FONT = "Arial"


# ---------------------------------------------------------------------------
# Low-level OOXML helpers
# ---------------------------------------------------------------------------

def _shade(el, fill_hex: str) -> None:
    """Apply a solid background fill to a cell or paragraph properties element."""
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pr.append(shd)


def _shade_cell(cell, fill_hex: str) -> None:
    _shade(cell._tc, fill_hex)


def _para_border(paragraph, edge: str, color: str, sz: int = 6, space: int = 4) -> None:
    """Draw a single rule on one edge ('top'/'bottom') of a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    e = OxmlElement(f"w:{edge}")
    e.set(qn("w:val"), "single")
    e.set(qn("w:sz"), str(sz))
    e.set(qn("w:space"), str(space))
    e.set(qn("w:color"), color)
    pbdr.append(e)


def _table_borders(table, color: str, sz: int = 4) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    table._tbl.tblPr.append(borders)


def _no_table_borders(table) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    table._tbl.tblPr.append(borders)


def _run(paragraph, text: str, *, size: float, color: RGBColor = _INK,
         bold: bool = False, italic: bool = False):
    r = paragraph.add_run(text)
    r.font.name = _FONT
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    r.italic = italic
    return r


def _spacing(paragraph, *, before: float = 0.0, after: float = 4.0,
             line: float | None = None) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line is not None:
        pf.line_spacing = line


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------

def _h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _spacing(p, before=6, after=3)
    _run(p, text, size=19, color=_INK, bold=True)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _spacing(p, before=2, after=4)
    _run(p, text, size=13, color=_ACCENT_DK, bold=True)


def _h3(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _spacing(p, before=8, after=4)
    _run(p, text, size=12, color=_INK, bold=True)
    _para_border(p, "bottom", _ACCENT_HEX, sz=8, space=3)


def _p(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _spacing(p, after=6, line=1.15)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, text, size=10, color=_INK)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _spacing(p, after=2, line=1.1)
        _run(p, item, size=10, color=_INK)


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _spacing(p, before=2, after=6)
    _run(p, text, size=8.5, color=_MUTED, italic=True)


def _spacer(doc: Document, mm: float) -> None:
    p = doc.add_paragraph()
    _spacing(p, after=mm * 2.0)


def _table(doc: Document, columns: list[str], rows: list[list[str]],
           numeric_from: int) -> None:
    tbl = doc.add_table(rows=1, cols=len(columns))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    _table_borders(tbl, _LINE, sz=4)

    def _fill_row(cells, values, *, bold: bool, fill: str | None) -> None:
        for i, (cell, val) in enumerate(zip(cells, values)):
            if fill:
                _shade_cell(cell, fill)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            _spacing(para, before=1.5, after=1.5, line=1.05)
            right = numeric_from is not None and i >= numeric_from and i != 0
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT if right else WD_ALIGN_PARAGRAPH.LEFT
            _run(para, str(val), size=9, color=_INK, bold=bold)

    _fill_row(tbl.rows[0].cells, columns, bold=True, fill=_HEAD_BG)
    for idx, row in enumerate(rows):
        cells = tbl.add_row().cells
        _fill_row(cells, row, bold=False, fill=_ZEBRA if idx % 2 == 0 else None)

    _spacer(doc, 1.5)


def _statcards(doc: Document, cards: list[dict]) -> None:
    """The 'Key Financial Highlights' panel: a single continuous cream block,
    laid out as a borderless 2-column grid (bold amber value + description)."""
    ncards = len(cards)
    nrows = (ncards + 1) // 2
    tbl = doc.add_table(rows=nrows, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    _no_table_borders(tbl)

    for idx, card in enumerate(cards):
        cell = tbl.rows[idx // 2].cells[idx % 2]
        _shade_cell(cell, _CREAM)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

        val_p = cell.paragraphs[0]
        _spacing(val_p, before=4, after=2)
        _run(val_p, card.get("value", ""), size=15, color=_AMBER, bold=True)

        txt_p = cell.add_paragraph()
        _spacing(txt_p, after=6, line=1.1)
        _run(txt_p, card.get("text", ""), size=9, color=_INK)

    # Any trailing empty cell (odd number of cards) still needs the cream fill
    # so the panel reads as one solid block.
    if ncards % 2 == 1:
        _shade_cell(tbl.rows[-1].cells[1], _CREAM)

    _spacer(doc, 2.0)


_RENDER = {
    "h1": lambda doc, b: _h1(doc, b["text"]),
    "h2": lambda doc, b: _h2(doc, b["text"]),
    "h3": lambda doc, b: _h3(doc, b["text"]),
    "p": lambda doc, b: _p(doc, b["text"]),
    "bullets": lambda doc, b: _bullets(doc, b["items"]),
    "note": lambda doc, b: _note(doc, b["text"]),
    "table": lambda doc, b: _table(doc, b["columns"], b["rows"], b.get("numeric_from", 1)),
    "statcards": lambda doc, b: _statcards(doc, b["cards"]),
    "spacer": lambda doc, b: _spacer(doc, b.get("mm", 3.0)),
}


# ---------------------------------------------------------------------------
# Page furniture: default styles, running header + footer, cover
# ---------------------------------------------------------------------------

def _configure(doc: Document, ds: CDADataset) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(10)
    normal.font.color.rgb = _INK

    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    for side in ("top", "bottom", "left", "right"):
        setattr(section, f"{side}_margin", Mm(18 if side in ("top", "bottom") else 16))
    # A clean cover: no running header/footer on the first page.
    section.different_first_page_header_footer = True

    # Running header (pages 2+).
    hp = section.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(hp, f"{ds.company} — Compensation Discussion & Analysis", size=8, color=_MUTED)
    _para_border(hp, "bottom", _LINE, sz=4, space=2)

    # Running footer (pages 2+): "<proxy_year> Proxy Statement   genpact".
    fp = section.footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _para_border(fp, "top", _LINE, sz=4, space=6)
    _run(fp, f"{ds.proxy_year} Proxy Statement   ", size=9, color=_MUTED)
    _run(fp, ds.company.lower(), size=10.5, color=_INK, bold=True)


def _cover(doc: Document, ds: CDADataset) -> None:
    doc.add_paragraph().paragraph_format.space_after = Pt(90)

    kicker = doc.add_paragraph()
    _spacing(kicker, after=6)
    _run(kicker, f"{ds.company.upper()}  ·  {ds.proxy_year} PROXY STATEMENT",
         size=10, color=_ACCENT, bold=True)

    title = doc.add_paragraph()
    _spacing(title, after=2)
    _run(title, "Compensation Discussion & Analysis", size=28, color=_INK, bold=True)

    sub = doc.add_paragraph()
    _spacing(sub, after=8)
    _run(sub, "Named Executive Officers — Excluding CEO", size=13, color=_MUTED)

    rule = doc.add_paragraph()
    _para_border(rule, "bottom", "C0392B", sz=16, space=1)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def build_docx(ds: CDADataset, blocks: list[Block]) -> bytes:
    doc = Document()
    _configure(doc, ds)
    _cover(doc, ds)
    doc.add_page_break()

    for block in blocks:
        renderer = _RENDER.get(block["type"])
        if renderer is not None:
            renderer(doc, block)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
